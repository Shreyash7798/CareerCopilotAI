"""Resume engine (spec section 12).

Takes the user's master resume (DOCX) plus a job description and produces a
tailored, ATS-friendly DOCX:

* Keeps every fact from the master resume — nothing is fabricated.
* Reorders the skills line so JD-matched skills come first.
* Reorders bullet points inside each experience block so the most relevant
  (highest keyword overlap with the JD) appear first.
* Adds a "Targeted for" line under the header naming the role/company.

PDF conversion is attempted with LibreOffice (`soffice`) if it is installed;
otherwise only the DOCX is produced.
"""

from __future__ import annotations

import json
import re
import shutil
import subprocess
from datetime import datetime
from pathlib import Path

from docx import Document

from app.config import data_dir

TAILOR_DOCX_NAME = "master_tailor.docx"

STOPWORDS = {
    "the", "and", "for", "with", "you", "your", "our", "are", "will", "have",
    "this", "that", "from", "their", "they", "them", "who", "what", "all",
    "can", "not", "but", "has", "was", "were", "been", "being", "its", "it's",
    "into", "than", "then", "when", "where", "which", "while", "would", "should",
    "could", "may", "might", "must", "shall", "about", "above", "after", "before",
    "job", "role", "work", "team", "company", "candidate", "years", "experience",
    "including", "ability", "strong", "skills", "required", "preferred", "etc",
}

WORD_RE = re.compile(r"[a-zA-Z][a-zA-Z+#./-]{2,}")


def extract_keywords(text: str, top_n: int = 40) -> list[str]:
    """Frequency-ranked content words from a job description."""
    counts: dict[str, int] = {}
    for word in WORD_RE.findall((text or "").lower()):
        if word in STOPWORDS:
            continue
        counts[word] = counts.get(word, 0) + 1
    ranked = sorted(counts.items(), key=lambda kv: (-kv[1], kv[0]))
    return [w for w, _ in ranked[:top_n]]


def _overlap_score(text: str, keywords: list[str]) -> int:
    lowered = text.lower()
    return sum(1 for k in keywords if k in lowered)


def _is_bullet(paragraph) -> bool:
    style = (paragraph.style.name or "").lower()
    text = paragraph.text.strip()
    return bool(text) and ("list" in style or "bullet" in style or text[:1] in "•-–▪●○*")


def _reorder_bullet_runs(doc: Document, keywords: list[str]) -> int:
    """Reorder consecutive bullet paragraphs by JD keyword overlap.

    Only the text is swapped between paragraphs, so formatting, styles and
    document structure are preserved (ATS-safe).
    """
    paragraphs = doc.paragraphs
    reordered = 0
    i = 0
    while i < len(paragraphs):
        if _is_bullet(paragraphs[i]):
            j = i
            while j < len(paragraphs) and _is_bullet(paragraphs[j]):
                j += 1
            block = paragraphs[i:j]
            if len(block) > 1:
                texts = [p.text for p in block]
                ranked = sorted(texts, key=lambda t: -_overlap_score(t, keywords))
                if ranked != texts:
                    for para, new_text in zip(block, ranked):
                        _set_paragraph_text(para, new_text)
                    reordered += 1
            i = j
        else:
            i += 1
    return reordered


def _set_paragraph_text(paragraph, text: str) -> None:
    """Replace paragraph text while keeping the first run's formatting."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _reorder_skill_lines(doc: Document, keywords: list[str]) -> int:
    """Reorder comma-separated items in lines under a Skills heading."""
    changed = 0
    in_skills = False
    for para in doc.paragraphs:
        text = para.text.strip()
        lowered = text.lower()
        if lowered and len(lowered) < 40 and "skill" in lowered and len(lowered.split()) <= 4:
            in_skills = True
            continue
        if in_skills:
            if not text:
                continue
            if len(text) < 40 and text.endswith(":") or (len(text.split()) <= 4 and "," not in text and not _is_bullet(para)):
                in_skills = False
                continue
            items = [s.strip() for s in text.split(",") if s.strip()]
            if len(items) > 2:
                ranked = sorted(items, key=lambda s: -_overlap_score(s, keywords))
                if ranked != items:
                    _set_paragraph_text(para, ", ".join(ranked))
                    changed += 1
    return changed


def _slug(value: str) -> str:
    return re.sub(r"[^a-zA-Z0-9]+", "-", value).strip("-").lower() or "job"


def _bullet_text(line: str) -> str | None:
    stripped = line.strip()
    if not stripped:
        return None
    if stripped[:1] in "•-*–▪●○":
        return stripped.lstrip("•-*–▪●○ ").strip()
    if re.match(r"^[-–]\s+", stripped):
        return re.sub(r"^[-–]\s+", "", stripped).strip()
    return None


def build_master_docx_from_text(raw_text: str, *, parsed: dict | None = None) -> Document:
    """Build a tailor-friendly DOCX from plain CV text (PDF/TXT uploads)."""
    parsed = parsed or {}
    doc = Document()
    if parsed.get("full_name"):
        doc.add_paragraph(str(parsed["full_name"]))
    contact = " · ".join(
        part for part in (parsed.get("email") or "", parsed.get("phone") or "") if part
    )
    if contact:
        doc.add_paragraph(contact)
    skills = parsed.get("skills") or []
    if skills:
        doc.add_paragraph("Skills")
        doc.add_paragraph(", ".join(skills))

    in_experience = False
    for line in raw_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        lowered = stripped.lower()
        if len(stripped) < 48 and any(
            h in lowered for h in ("experience", "employment", "work history")
        ):
            doc.add_heading(stripped, level=2)
            in_experience = True
            continue
        bullet = _bullet_text(stripped)
        if bullet:
            doc.add_paragraph(bullet, style="List Bullet")
            in_experience = True
            continue
        doc.add_paragraph(stripped)
    return doc


def write_tailor_master_docx(raw_text: str, *, parsed: dict | None = None) -> Path:
    out_dir = data_dir() / "cv"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / TAILOR_DOCX_NAME
    build_master_docx_from_text(raw_text, parsed=parsed).save(str(out_path))
    return out_path


def convert_doc_to_docx(source: Path) -> Path | None:
    """Convert legacy .doc to .docx when LibreOffice is available."""
    if source.suffix.lower() != ".doc":
        return None
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    out_dir = source.parent
    try:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "docx", "--outdir", str(out_dir), str(source)],
            check=True,
            capture_output=True,
            timeout=120,
        )
    except (subprocess.SubprocessError, OSError):
        return None
    out = source.with_suffix(".docx")
    return out if out.exists() else None


def resolve_master_docx_path(
    cv_path: str | None,
    *,
    profile_json: str | None = None,
) -> Path | None:
    """Return a DOCX master resume path suitable for tailoring."""
    parsed: dict = {}
    if profile_json:
        try:
            loaded = json.loads(profile_json)
            if isinstance(loaded, dict):
                parsed = loaded
        except json.JSONDecodeError:
            parsed = {}

    if cv_path:
        path = Path(cv_path)
        if path.suffix.lower() == ".docx" and path.exists():
            return path
        if path.suffix.lower() == ".doc" and path.exists():
            converted = convert_doc_to_docx(path)
            if converted and converted.exists():
                return converted

    tailor_path = data_dir() / "cv" / TAILOR_DOCX_NAME
    if tailor_path.exists():
        return tailor_path

    raw_text = parsed.get("raw_text") or ""
    if not raw_text and cv_path:
        source = Path(cv_path)
        if source.exists():
            try:
                from app.cv_parser import extract_text

                raw_text = extract_text(source)
            except (ValueError, OSError):
                raw_text = ""

    if not raw_text.strip():
        return None

    return write_tailor_master_docx(raw_text, parsed=parsed)


def _add_targeted_for_line(doc: Document, job_title: str, company: str) -> None:
    label = f"Targeted for: {job_title} — {company}".strip(" —")
    if not doc.paragraphs:
        doc.add_paragraph(label)
        return
    targeted = doc.paragraphs[0].insert_paragraph_before(label)
    if targeted.runs:
        targeted.runs[0].italic = True


def convert_to_pdf(docx_path: Path) -> Path | None:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    out_dir = docx_path.parent
    try:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)],
            check=True,
            capture_output=True,
            timeout=120,
        )
    except (subprocess.SubprocessError, OSError):
        return None
    pdf = docx_path.with_suffix(".pdf")
    return pdf if pdf.exists() else None


def tailor_resume(
    master_path: str | Path,
    *,
    job_title: str,
    company: str,
    job_description: str,
) -> dict:
    """Produce a tailored DOCX (and PDF when possible) for one job.

    Returns {"docx": path, "pdf": path|None, "matched_keywords": [...]}.
    """
    master_path = Path(master_path)
    if master_path.suffix.lower() != ".docx":
        raise ValueError("The master resume must be a .docx file for tailoring.")

    jd_text = (job_description or "").strip()
    if not jd_text:
        jd_text = f"{job_title} at {company}"
    keywords = extract_keywords(jd_text)
    doc = Document(str(master_path))

    resume_text = "\n".join(p.text for p in doc.paragraphs).lower()
    matched = [k for k in keywords if k in resume_text]

    _add_targeted_for_line(doc, job_title, company)
    _reorder_bullet_runs(doc, matched or keywords)
    _reorder_skill_lines(doc, matched or keywords)

    out_dir = data_dir() / "resumes"
    out_dir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    out_path = out_dir / f"{_slug(company)}-{_slug(job_title)}-{stamp}.docx"
    doc.save(str(out_path))

    pdf_path = convert_to_pdf(out_path)
    return {
        "docx": str(out_path),
        "pdf": str(pdf_path) if pdf_path else None,
        "matched_keywords": matched,
        "matched_keywords_json": json.dumps(matched),
    }
