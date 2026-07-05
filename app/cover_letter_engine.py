"""Cover letter engine — local, zero-cost, no fabrication.

Builds a tailored cover letter DOCX from the user's profile/CV facts and
the job description. Only references experience and skills that appear in the
uploaded profile; it never invents employers, dates or credentials.
"""

from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt

from app.config import data_dir, get_profile
from app.resume_engine import extract_keywords, _slug


def _profile_facts(profile: dict, cv_json: dict | None) -> dict:
    merged = dict(profile)
    if cv_json:
        for key in ("full_name", "email", "phone", "experience_years", "skills", "employers"):
            if cv_json.get(key) and not merged.get(key):
                merged[key] = cv_json[key]
    return merged


def _matched_skills(profile_skills: list[str], jd_keywords: list[str]) -> list[str]:
    if not profile_skills:
        return []
    jd_set = set(jd_keywords)
    hits = []
    for skill in profile_skills:
        tokens = [t.lower() for t in re.findall(r"[a-zA-Z+#]+", skill.lower()) if len(t) > 2]
        if any(t in jd_set for t in tokens) or any(t in skill.lower() for t in jd_keywords[:15]):
            hits.append(skill)
    return hits[:8]


def _experience_sentence(profile: dict) -> str:
    years = profile.get("experience_years")
    employer = profile.get("current_employer") or ""
    employers = profile.get("employers") or profile.get("previous_employers") or []
    if isinstance(employers, str):
        employers = [employers]
    parts = []
    if years:
        parts.append(f"{years:g} years of professional experience")
    if employer:
        parts.append(f"currently at {employer}")
    elif employers:
        parts.append(f"with experience at {employers[0]}")
    domains = profile.get("preferred_domains") or profile.get("interests") or []
    if domains:
        parts.append(f"focused on {', '.join(domains[:3])}")
    return ", ".join(parts) + "." if parts else ""


def generate_cover_letter(
    *,
    job_title: str,
    company: str,
    job_description: str,
    profile: dict | None = None,
    cv_json: dict | None = None,
) -> dict:
    """Return {docx, matched_keywords, text}."""
    profile = _profile_facts(profile or get_profile(), cv_json)
    name = profile.get("full_name") or "Applicant"
    email = profile.get("email") or ""
    phone = profile.get("phone") or ""
    jd_keywords = extract_keywords(job_description, top_n=30)
    skills = _matched_skills(profile.get("skills") or [], jd_keywords)
    exp_sentence = _experience_sentence(profile)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    header = doc.add_paragraph()
    header.add_run(name).bold = True
    if email or phone:
        doc.add_paragraph(" · ".join(x for x in (email, phone) if x))
    doc.add_paragraph()

    doc.add_paragraph(datetime.now().strftime("%d %B %Y"))
    doc.add_paragraph()
    doc.add_paragraph(f"Re: {job_title} — {company}")
    doc.add_paragraph()

    doc.add_paragraph(
        f"Dear Hiring Team at {company},"
    )
    doc.add_paragraph(
        f"I am writing to express my interest in the {job_title} role. "
        f"Your posting emphasises {', '.join(jd_keywords[:5]) if jd_keywords else 'the responsibilities described'}, "
        f"which aligns closely with my background."
    )
    if exp_sentence:
        doc.add_paragraph(f"I bring {exp_sentence}")
    if skills:
        doc.add_paragraph(
            "Relevant strengths from my background include "
            + ", ".join(skills)
            + ", each of which maps directly to the requirements in your job description."
        )
    interests = profile.get("interests") or profile.get("preferred_domains") or []
    if interests:
        doc.add_paragraph(
            f"I am particularly motivated by work in {', '.join(interests[:4])}, "
            f"and would welcome the opportunity to contribute to {company}'s initiatives in this space."
        )
    doc.add_paragraph(
        "I would appreciate the opportunity to discuss how my experience can support your team. "
        "Thank you for your time and consideration."
    )
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph(name)

    out_dir = data_dir() / "cover_letters"
    out_dir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    out_path = out_dir / f"{_slug(company)}-{_slug(job_title)}-{stamp}.docx"
    doc.save(str(out_path))

    plain = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return {
        "docx": str(out_path),
        "matched_keywords": skills,
        "matched_keywords_json": json.dumps(skills),
        "text": plain,
    }
