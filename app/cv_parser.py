"""CV parsing: turn an uploaded PDF/DOCX resume into a structured profile.

Deliberately heuristic and fully local (no external APIs). The result
pre-fills the profile in settings.yaml; the user can then refine it in the
Settings screen. The CV remains the source of truth for personalization.
"""

from __future__ import annotations

import re
from pathlib import Path

EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(r"(?:\+?\d{1,3}[\s-]?)?(?:\(?\d{2,5}\)?[\s-]?)?\d{3,5}[\s-]?\d{4,6}")
YEARS_RE = re.compile(r"(\d{1,2})(?:\.\d)?\s*\+?\s*years?", re.IGNORECASE)
DATE_RANGE_RE = re.compile(
    r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*(\d{4})\s*[-–—to]+\s*"
    r"(?:(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*(\d{4})|present|current)",
    re.IGNORECASE,
)

# Skill vocabulary used for extraction. Extend freely; this is a starting
# point biased towards consulting/business roles plus common technical skills.
SKILL_VOCABULARY = [
    "Business Transformation", "Operations Consulting", "Manufacturing", "Supply Chain",
    "Strategy", "Process Improvement", "Process Excellence", "Lean", "Six Sigma", "Kaizen",
    "Stakeholder Management", "Project Management", "Program Management", "Change Management",
    "Market Research", "Due Diligence", "Cost Optimization", "Procurement", "Logistics",
    "Inventory Management", "Demand Planning", "S&OP", "ERP", "SAP", "Oracle",
    "Data Analysis", "Data Analytics", "Financial Modeling", "Business Case",
    "Excel", "PowerPoint", "Power BI", "Tableau", "SQL", "Python", "Alteryx",
    "Consulting", "Advisory", "Client Management", "Business Development",
    "Agile", "Scrum", "PMO", "KPI", "Benchmarking", "Root Cause Analysis",
]

SECTION_HEADERS = [
    "experience", "work experience", "professional experience", "employment",
    "education", "skills", "projects", "certifications", "summary", "profile",
]


def extract_text(path: str | Path) -> str:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix in (".docx", ".doc"):
        return _extract_docx(path)
    if suffix in (".txt", ".md"):
        return path.read_text(encoding="utf-8", errors="ignore")
    raise ValueError(f"Unsupported CV format '{suffix}'. Use PDF, DOCX or TXT.")


def _extract_pdf(path: Path) -> str:
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
    return "\n".join(parts)


def _extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _guess_name(lines: list[str]) -> str:
    for line in lines[:8]:
        candidate = line.strip()
        if not candidate or EMAIL_RE.search(candidate) or PHONE_RE.fullmatch(candidate):
            continue
        words = candidate.split()
        if 1 < len(words) <= 5 and all(w[:1].isupper() or not w[:1].isalpha() for w in words):
            lowered = candidate.lower()
            if not any(h in lowered for h in SECTION_HEADERS) and "resume" not in lowered:
                return candidate
    return ""


def _guess_experience_years(text: str) -> float | None:
    lowered = text.lower()
    m = re.search(r"(\d{1,2}(?:\.\d)?)\s*\+?\s*years?(?:\s+of)?\s+(?:experience|exp)", lowered)
    if m:
        return float(m.group(1))
    # Fall back to summing date ranges in the experience section.
    ranges = DATE_RANGE_RE.findall(text)
    if ranges:
        from datetime import datetime

        current_year = datetime.now().year
        total = 0
        for start, end in ranges:
            end_year = int(end) if end else current_year
            total += max(0, end_year - int(start))
        if total:
            return float(min(total, 50))
    return None


def _extract_skills(text: str) -> list[str]:
    lowered = text.lower()
    return [s for s in SKILL_VOCABULARY if s.lower() in lowered]


def _extract_employers(text: str) -> list[str]:
    """Best-effort: lines near date ranges in the experience section.

    CV layouts differ ("Company / Title / Dates" vs "Title, Company | Dates"),
    so all nearby candidate lines are kept and the user confirms in Settings.
    """
    employers: list[str] = []
    lines = text.splitlines()
    for i, line in enumerate(lines):
        if DATE_RANGE_RE.search(line):
            for j in (i, i - 1, i - 2):
                if 0 <= j < len(lines):
                    candidate = lines[j].strip()
                    candidate = DATE_RANGE_RE.sub("", candidate).strip(" -–—|,")
                    if (
                        candidate
                        and len(candidate) < 80
                        and not EMAIL_RE.search(candidate)
                        and candidate.lower() not in SECTION_HEADERS
                        and candidate not in employers
                    ):
                        employers.append(candidate)
    return employers[:8]


def parse_cv(path: str | Path) -> dict:
    """Parse a CV file into a structured profile dict."""
    text = extract_text(path)
    lines = [ln for ln in text.splitlines() if ln.strip()]

    email_match = EMAIL_RE.search(text)
    phone_match = PHONE_RE.search(text.replace("\n", " "))
    years = _guess_experience_years(text)
    employers = _extract_employers(text)

    return {
        "full_name": _guess_name(lines),
        "email": email_match.group(0) if email_match else "",
        "phone": phone_match.group(0).strip() if phone_match else "",
        "experience_years": years,
        "skills": _extract_skills(text),
        "employers": employers,
        "raw_text": text,
    }
