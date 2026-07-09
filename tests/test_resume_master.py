import json
from pathlib import Path

from docx import Document

from app.resume_engine import resolve_master_docx_path, write_tailor_master_docx


def test_resolve_master_docx_from_pdf_profile(tmp_path, monkeypatch):
    import app.resume_engine as re_mod

    monkeypatch.setattr(re_mod, "data_dir", lambda: tmp_path)

    raw = """Rahul Sharma
rahul@example.com

Skills
Excel, Supply Chain, Strategy

Experience
PwC India — Associate
- Optimized supply chain networks
- Built PowerPoint decks for clients
"""
    parsed = {
        "full_name": "Rahul Sharma",
        "email": "rahul@example.com",
        "skills": ["Supply Chain", "Excel"],
        "raw_text": raw,
    }
    pdf_path = tmp_path / "cv" / "master.pdf"
    pdf_path.parent.mkdir(parents=True)
    pdf_path.write_bytes(b"%PDF-1.4")

    resolved = resolve_master_docx_path(str(pdf_path), profile_json=json.dumps(parsed))
    assert resolved is not None
    assert resolved.suffix.lower() == ".docx"
    assert resolved.exists()

    doc = Document(str(resolved))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Rahul Sharma" in text
    assert "supply chain" in text.lower()


def test_write_tailor_master_docx_creates_bullets(tmp_path, monkeypatch):
    import app.resume_engine as re_mod

    monkeypatch.setattr(re_mod, "data_dir", lambda: tmp_path)
    path = write_tailor_master_docx(
        "Experience\n- Led supply chain projects\n- Excel reporting",
        parsed={"full_name": "Test User"},
    )
    doc = Document(str(path))
    bullets = [p.text for p in doc.paragraphs if "supply chain" in p.text.lower()]
    assert bullets
