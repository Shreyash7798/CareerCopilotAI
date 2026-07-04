from docx import Document

from app.resume_engine import extract_keywords, tailor_resume


def _master(tmp_path):
    doc = Document()
    doc.add_paragraph("Rahul Sharma")
    doc.add_paragraph("Skills")
    doc.add_paragraph("PowerPoint, Excel, Supply Chain, Strategy")
    doc.add_heading("Experience", level=2)
    doc.add_paragraph("PwC India — Senior Associate")
    doc.add_paragraph("- Prepared client presentations in PowerPoint", style="List Bullet")
    doc.add_paragraph("- Optimized supply chain networks reducing cost by 12%", style="List Bullet")
    path = tmp_path / "master.docx"
    doc.save(str(path))
    return path


def test_extract_keywords():
    kws = extract_keywords("We need supply chain experience. Supply chain and Excel skills.")
    assert "supply" in kws and "chain" in kws and "excel" in kws
    assert "the" not in kws


def test_tailor_resume_reorders_without_fabricating(tmp_path, monkeypatch):
    import app.resume_engine as re_mod

    monkeypatch.setattr(re_mod, "data_dir", lambda: tmp_path)
    master = _master(tmp_path)
    jd = "Looking for a supply chain consultant. Strong supply chain and network optimization."
    result = tailor_resume(master, job_title="Supply Chain Consultant", company="Acme", job_description=jd)

    out = Document(result["docx"])
    texts = [p.text for p in out.paragraphs]
    bullets = [t for t in texts if "PowerPoint" in t or "supply chain" in t.lower()]
    # supply-chain bullet should now be ranked above the PowerPoint bullet
    sc_idx = next(i for i, t in enumerate(texts) if "supply chain networks" in t.lower())
    ppt_idx = next(i for i, t in enumerate(texts) if "client presentations" in t)
    assert sc_idx < ppt_idx
    # no fabricated content: every non-empty paragraph existed in the master
    original = {p.text for p in Document(str(master)).paragraphs}
    reordered_skills = {"Supply Chain, Strategy, PowerPoint, Excel"}
    for t in texts:
        if not t.strip():
            continue
        assert t in original or set(t.split(", ")) == {"PowerPoint", "Excel", "Supply Chain", "Strategy"}
    assert "supply" in result["matched_keywords"]
